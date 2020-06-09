from docx import Document
import datetime as dt
from docx.shared import Pt
from docx2pdf import convert
import shutil
import os
document= Document('CL_Template.docx')
today= dt.date.today()


finance=['Advanced Corporate Finance,',' Financial Statement Analysis and Valuations,', ' and Big Data Analytics']
marketing=['Consumer Behavior,',' Strategic Marketing Communications,', ' and Big Data and Marketing Analytics']
economics=['International Macroeconomics,',' International Trade,', ' and International Finance']
data=['Big Data Analytics,', ' Financial Trading Strategies,', ' and Quantitative Methods in Economics']
ib=['Advanced Corporate Finance,',' Investments,',' and Capital Market Theory']
trading=['Financial Trading Strategies,', ' and Investments']

def paragraph_spacing(document):
    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.space_before = Pt(.25)
    paragraph_format.space_after = Pt(.25)

def to_string(s):  
    str1 = ""     
    for ele in s:  
        str1 += ele     
    return str1 

def cover_letter(document, position, company, classes, job_description):
    document.add_paragraph(today.strftime("%B %d, %Y"))
    paragraph_spacing(document)
    document.add_paragraph('Dear Hiring Manager:')
    paragraph_spacing(document)
    document.add_paragraph('I am writing to express my interest in the position of ' + position +  ' at ' + company+ ', where my academic training and professional experience will enable me to make a positive impact on business growth and success, while at the same time allowing me to learn and grow.')
    paragraph_spacing(document)
    document.add_paragraph('I have completed my education and received my Bachelor of Commerce degree from the University in Toronto in June of 2020, as a Management Specialist with a Finance focus and a major in Economics. During my time at the university, I have completed a wide range of business courses as well as international and domestic internships with Sobeys, Microsoft, and Antalis.')
    paragraph_spacing(document)
    document.add_paragraph('My time at Sobeys last summer was very valuable to me. The family-like environment I was welcomed to allowed me to quickly adapt to my new position. My first main responsibility at Sobeys was entering new product details on to SAP, which helped me get familiar with the system and learn more about how Sobeys establishes communication channels with suppliers and warehouses. My second main responsibility was analyzing sales data of deli meat items and finding correlations between weekly sales data and the corresponding promotions. I learned the importance of different factors that goes into pricing and promotions of these items, such as shrink, or other products that customers purchase with their deli meats. I had the opportunity to prepare promotional strategies for these items based on margin analysis.')
    paragraph_spacing(document)
    document.add_paragraph('My previous internships have taught me many valuable skills and experiences. During my time at Microsoft, I learned forecasting and analyzing sales data based on the results of previous periods. Furthermore, I saved $2,000 per month in lodging and travel costs by developing more efficient store visit schedules for field sales consultants, based on store locations. ')
    paragraph_spacing(document)
    document.add_paragraph('Throughout my studies, I was enrolled in classes that further appended on my learnings from my previous internships. Classes such as '+ classes + ' gave me new perspectives and skills that I can use for '+ job_description +'. Furthermore, I am familiar with using tools such as Power BI, Tableau, and SAP, and performing qualitative and quantitative analyses using Python and Excel. In addition, I am currently pursuing the Certificate in Big Data Analytics from York University and I am a candidate for CFA Level 1 December exam. ')
    paragraph_spacing(document)
    document.add_paragraph("I'm very enthusiastic about the possibility of joining "+company+ ". If you need additional information, or you would like to discuss in person the skills I have learned that would be very useful in the role of "+position+ ", I'd be happy to meet with you. Thank you for your time and consideration.")
    paragraph_spacing(document)
    document.add_paragraph('Sincerely yours,')
    document.add_paragraph('Burak Nevzat Yalcin')
    document.save('Cover_Letter.docx')
    convert('Cover_Letter.docx','Burak Nevzat Yalcin Cover Letter.pdf')
    shutil.move(os.path.join('C:\\Users\\Owner\\Documents\\Burak Nevzat Yalcin Cover Letter.pdf'),os.path.join('C:\\Users\\Owner\\Desktop\\Documents\\Others\\Resume and Letter\\New Docs\\Burak Nevzat Yalcin Cover Letter.pdf'))